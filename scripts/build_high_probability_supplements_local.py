from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt


PROJECT_ROOT = Path(
    r"R:\NAS_DRIVE\IMUT\1-Research_Output\1-Papers\1_In_Preparation\2026-Abaqus-experimentclarify-archive"
)
AUTOBOOST = PROJECT_ROOT / "ai_autoboost"
RAW_SEPH_DIR = Path(
    r"R:\NAS_DRIVE\IMUT\1-Research_Output\7-code\1-博士论文相关代码\1-Abaqus自动化建模代码\2-分离单独\isight文件\Isight_Catia_Abaqus\Abaqus_dataTEsepH"
)
SOFTWAREX_ROOT = Path(
    r"R:\NAS_DRIVE\IMUT\1-Research_Output\1-Papers\2-已发表\21-SoftwareX\4-Reivsed\shearcode"
)
SOFTWAREX_WORK = SOFTWAREX_ROOT / "work"
V7_MANUSCRIPT = AUTOBOOST / "gpt_interaction_lift" / "manuscript" / "manuscript_v7_final_gpt_reviewed.md"

SUP_ROOT = AUTOBOOST / "high_probability_supplements"
OUT_DIR = SUP_ROOT / "outputs"
DOC_DIR = SUP_ROOT / "docs"
MANUSCRIPT_DIR = SUP_ROOT / "manuscript"
FIG_DIR = SUP_ROOT / "figures"
PACKAGE_DIR = AUTOBOOST / "submission_JCSR" / "package_v8_high_probability_supplemented"

NA = "not_available"


@dataclass
class StaSummary:
    exists: bool
    completed: bool | None
    final_line: str
    final_step: str
    final_increment: str
    final_attempt: str
    final_total_iters: str
    final_step_time_lpf: str
    final_inc_time_lpf: str
    row_count: int


def ensure_dirs() -> None:
    for path in [OUT_DIR, DOC_DIR, MANUSCRIPT_DIR, FIG_DIR, PACKAGE_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def read_text(path: Path) -> str:
    if not path.exists():
        return ""
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            return path.read_text(encoding=enc, errors="replace")
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")


def rel_or_abs(path: Path) -> str:
    try:
        return str(path.relative_to(PROJECT_ROOT))
    except ValueError:
        return str(path)


def sha256_head(path: Path, max_bytes: int = 1024 * 1024) -> str:
    if not path.exists() or not path.is_file():
        return ""
    h = hashlib.sha256()
    with path.open("rb") as fh:
        remaining = max_bytes
        while remaining > 0:
            chunk = fh.read(min(65536, remaining))
            if not chunk:
                break
            h.update(chunk)
            remaining -= len(chunk)
    return h.hexdigest()


def file_inventory(root: Path) -> pd.DataFrame:
    rows = []
    for p in sorted(root.glob("*")):
        if not p.is_file():
            continue
        stat = p.stat()
        rows.append(
            {
                "file_name": p.name,
                "extension": p.suffix.lower(),
                "absolute_path": str(p),
                "size_bytes": stat.st_size,
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
                "sha256_first_1mb": sha256_head(p),
            }
        )
    return pd.DataFrame(rows)


def parse_sta(path: Path) -> StaSummary:
    if not path.exists():
        return StaSummary(False, None, "", "", "", "", "", "", "", 0)
    text = read_text(path)
    completed = "THE ANALYSIS HAS COMPLETED SUCCESSFULLY" in text
    if "THE ANALYSIS HAS NOT BEEN COMPLETED" in text:
        completed = False
    numeric_rows = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if re.match(r"^\s*\d+\s+\d+\s+\S+", line):
            tokens = line.split()
            if len(tokens) >= 8:
                numeric_rows.append((line, tokens))
    if not numeric_rows:
        return StaSummary(True, completed, "", "", "", "", "", "", "", 0)
    final_line, tok = numeric_rows[-1]
    return StaSummary(
        True,
        completed,
        final_line.strip(),
        tok[0],
        tok[1],
        tok[2],
        tok[5] if len(tok) > 5 else "",
        tok[6] if len(tok) > 6 else "",
        tok[7] if len(tok) > 7 else "",
        len(numeric_rows),
    )


def parse_msg(path: Path) -> dict:
    text = read_text(path)
    if not text:
        return {
            "msg_exists": False,
            "analysis_terminated_due_to_errors": NA,
            "minimum_time_increment_error": NA,
            "solution_diverging_note": NA,
            "total_increments": NA,
            "cutbacks": NA,
            "iterations": NA,
            "negative_eigenvalue_warnings": NA,
            "error_messages": NA,
            "warning_messages_analysis": NA,
            "wallclock_seconds": NA,
            "processors_reported": NA,
        }
    def find_int(pattern: str) -> str:
        m = re.search(pattern, text, flags=re.IGNORECASE)
        return m.group(1).strip() if m else NA
    def find_last_int(pattern: str) -> str:
        hits = re.findall(pattern, text, flags=re.IGNORECASE)
        return hits[-1].strip() if hits else NA

    return {
        "msg_exists": True,
        "analysis_terminated_due_to_errors": "THE ANALYSIS HAS BEEN TERMINATED DUE TO PREVIOUS ERRORS" in text,
        "minimum_time_increment_error": "TIME INCREMENT REQUIRED IS LESS THAN THE MINIMUM SPECIFIED" in text,
        "solution_diverging_note": "THE SOLUTION APPEARS TO BE DIVERGING" in text,
        "total_increments": find_int(r"TOTAL OF\s+(\d+)\s+INCREMENTS"),
        "cutbacks": find_int(r"\n\s*(\d+)\s+CUTBACKS IN AUTOMATIC INCREMENTATION"),
        "iterations": find_int(r"\n\s*(\d+)\s+ITERATIONS INCLUDING CONTACT ITERATIONS"),
        "negative_eigenvalue_warnings": find_int(r"\n\s*(\d+)\s+ANALYSIS WARNINGS ARE NEGATIVE EIGENVALUE"),
        "error_messages": find_int(r"\n\s*(\d+)\s+ERROR MESSAGES"),
        "warning_messages_analysis": find_int(r"\n\s*(\d+)\s+WARNING MESSAGES DURING ANALYSIS"),
        "wallclock_seconds": find_last_int(r"WALLCLOCK TIME \(SEC\)\s*=\s*(\d+)"),
        "processors_reported": find_int(r"USING THE DIRECT SOLVER WITH\s+(\d+)\s+PROCESSORS"),
    }


def parse_inp(path: Path) -> dict:
    text = read_text(path)
    if not text:
        return {
            "inp_exists": False,
            "element_types": NA,
            "has_shell_section": NA,
            "has_static_riks": NA,
            "has_nlgeom_step": NA,
            "has_contact": NA,
            "has_friction": NA,
            "has_plastic_material": NA,
            "has_boundary": NA,
            "output_requests": NA,
        }
    element_types = sorted(set(re.findall(r"^\*Element,\s*type=([A-Za-z0-9_]+)", text, flags=re.MULTILINE | re.IGNORECASE)))
    output_reqs = sorted(set(re.findall(r"^\*(?:Node|Element|Contact) Output[^\n]*", text, flags=re.MULTILINE | re.IGNORECASE)))
    return {
        "inp_exists": True,
        "element_types": ";".join(element_types) if element_types else NA,
        "has_shell_section": "*Shell Section" in text,
        "has_static_riks": bool(re.search(r"^\*Static,\s*riks", text, flags=re.MULTILINE | re.IGNORECASE)),
        "has_nlgeom_step": bool(re.search(r"^\*Step,[^\n]*nlgeom\s*=\s*YES", text, flags=re.MULTILINE | re.IGNORECASE)),
        "has_contact": "*Contact" in text or "*Surface Interaction" in text,
        "has_friction": "*Friction" in text,
        "has_plastic_material": "*Plastic" in text,
        "has_boundary": "*Boundary" in text,
        "output_requests": "; ".join(output_reqs) if output_reqs else NA,
    }


def parse_rpt_result2(path: Path) -> dict:
    text = read_text(path)
    if not text:
        return {"result2_exists": False, "max_rf2": NA, "max_rf2_at_x": NA, "max_rf3": NA, "max_rf3_at_x": NA}
    # Abaqus report is column-formatted; extract the known MAXIMUM row and AT X row.
    max_line = ""
    at_line = ""
    for line in text.splitlines():
        if line.strip().startswith("MAXIMUM"):
            max_line = line
        if line.strip().startswith("AT X =") and max_line:
            at_line = line
            break
    nums_max = re.findall(r"[-+]?\d+(?:\.\d+)?(?:E[+-]?\d+)?", max_line)
    nums_at = re.findall(r"[-+]?\d+(?:\.\d+)?(?:E[+-]?\d+)?", at_line)
    return {
        "result2_exists": True,
        "max_rf2": nums_max[1] if len(nums_max) > 1 else NA,
        "max_rf2_at_x": nums_at[1] if len(nums_at) > 1 else NA,
        "max_rf3": nums_max[2] if len(nums_max) > 2 else NA,
        "max_rf3_at_x": nums_at[2] if len(nums_at) > 2 else NA,
    }


def raw_solver_audit() -> tuple[pd.DataFrame, pd.DataFrame]:
    inventory = file_inventory(RAW_SEPH_DIR)
    rows = []
    for job in ["Job-1", "Job-2"]:
        sta = parse_sta(RAW_SEPH_DIR / f"{job}.sta")
        msg = parse_msg(RAW_SEPH_DIR / f"{job}.msg")
        inp = parse_inp(RAW_SEPH_DIR / f"{job}.inp")
        rpt = parse_rpt_result2(RAW_SEPH_DIR / "Result2.rpt") if job == "Job-2" else {}
        log_text = read_text(RAW_SEPH_DIR / f"{job}.log")
        odb_path = RAW_SEPH_DIR / f"{job}.odb"
        rows.append(
            {
                "job": job,
                "source_directory": str(RAW_SEPH_DIR),
                "odb_exists": odb_path.exists(),
                "odb_size_bytes": odb_path.stat().st_size if odb_path.exists() else "",
                "sta_exists": sta.exists,
                "sta_completed_successfully": sta.completed,
                "sta_row_count": sta.row_count,
                "sta_final_line": sta.final_line,
                "sta_final_step": sta.final_step,
                "sta_final_increment": sta.final_increment,
                "sta_final_attempt": sta.final_attempt,
                "sta_final_total_iters": sta.final_total_iters,
                "sta_final_step_time_lpf": sta.final_step_time_lpf,
                "sta_final_inc_time_lpf": sta.final_inc_time_lpf,
                "log_exited_with_errors": "exited with errors" in log_text.lower(),
                **msg,
                **inp,
                **rpt,
            }
        )
    return inventory, pd.DataFrame(rows)


def iter_softwarex_cases() -> Iterable[Path]:
    if not SOFTWAREX_WORK.exists():
        return []
    return sorted(p for p in SOFTWAREX_WORK.glob("*") if p.is_dir())


def softwarex_solver_audit() -> pd.DataFrame:
    rows = []
    for case_dir in iter_softwarex_cases():
        for load_dir in sorted(case_dir.glob("*")):
            if not load_dir.is_dir():
                continue
            riks_sta = next(load_dir.glob("Job_Riks*.sta"), None)
            riks_msg = next(load_dir.glob("Job_Riks*.msg"), None)
            riks_inp = next(load_dir.glob("Job_Riks*.inp"), None)
            buckle_sta = next(load_dir.glob("Job_Buckle*.sta"), None)
            buckle_msg = next(load_dir.glob("Job_Buckle*.msg"), None)
            if not any([riks_sta, buckle_sta]):
                continue
            sta = parse_sta(riks_sta) if riks_sta else StaSummary(False, None, "", "", "", "", "", "", "", 0)
            msg = parse_msg(riks_msg) if riks_msg else {}
            inp = parse_inp(riks_inp) if riks_inp else {}
            b_sta = parse_sta(buckle_sta) if buckle_sta else StaSummary(False, None, "", "", "", "", "", "", "", 0)
            b_msg = parse_msg(buckle_msg) if buckle_msg else {}
            rows.append(
                {
                    "case_id": case_dir.name,
                    "load_case": load_dir.name,
                    "source_directory": str(load_dir),
                    "riks_sta_exists": bool(riks_sta),
                    "riks_sta_completed_successfully": sta.completed,
                    "riks_sta_row_count": sta.row_count,
                    "riks_final_step_time_lpf": sta.final_step_time_lpf,
                    "riks_final_inc_time_lpf": sta.final_inc_time_lpf,
                    "riks_msg_error_messages": msg.get("error_messages", NA),
                    "riks_msg_wallclock_seconds": msg.get("wallclock_seconds", NA),
                    "riks_element_types": inp.get("element_types", NA),
                    "riks_has_static_riks": inp.get("has_static_riks", NA),
                    "riks_has_contact": inp.get("has_contact", NA),
                    "riks_has_friction": inp.get("has_friction", NA),
                    "buckle_sta_exists": bool(buckle_sta),
                    "buckle_sta_completed_successfully": b_sta.completed,
                    "buckle_msg_error_messages": b_msg.get("error_messages", NA),
                    "riks_sta_path": str(riks_sta) if riks_sta else "",
                    "riks_inp_path": str(riks_inp) if riks_inp else "",
                }
            )
    return pd.DataFrame(rows)


def softwarex_validation_audit() -> pd.DataFrame:
    summary = SOFTWAREX_ROOT / "G1-G4_完整分析结果汇总.md"
    text = read_text(summary)
    rows = []
    # Expected table rows contain G1-G4, simulated kN, experimental kN, error %, LPF and displacement.
    pattern = re.compile(
        r"\|\s*\*\*(G\d)\*\*\s*\|[^|]*\|\s*\*\*?([0-9.]+)\*\*?\s*\|\s*([0-9.]+)\s*\|\s*\*\*?([+-]?[0-9.]+)%\*\*?\s*\|\s*([0-9.]+)\s*\|\s*([0-9.]+)\s*\|"
    )
    for m in pattern.finditer(text):
        rows.append(
            {
                "specimen": m.group(1),
                "fe_predicted_shear_kN": float(m.group(2)),
                "experimental_shear_kN": float(m.group(3)),
                "signed_error_percent": float(m.group(4)),
                "abs_error_percent": abs(float(m.group(4))),
                "reported_lpf": float(m.group(5)),
                "reported_max_displacement_mm": float(m.group(6)),
                "source_file": str(summary),
                "source_interpretation": "SoftwareX/Scandella shear-workflow validation summary; not a direct validation of P1 paired sensitivity ratios.",
            }
        )
    return pd.DataFrame(rows)


def make_softwarex_plot(df: pd.DataFrame) -> Path:
    out = FIG_DIR / "softwarex_solver_status_summary.png"
    if df.empty:
        return out
    status = df["riks_sta_completed_successfully"].fillna(False).map({True: "completed", False: "not completed"}).value_counts()
    fig, ax = plt.subplots(figsize=(5.2, 3.2))
    colors = ["#3B6EA8" if idx == "completed" else "#B85C38" for idx in status.index]
    ax.bar(status.index, status.values, color=colors)
    ax.set_ylabel("Number of Riks jobs")
    ax.set_title("SoftwareX shear benchmark solver status")
    ax.grid(axis="y", alpha=0.25)
    for i, v in enumerate(status.values):
        ax.text(i, v + 0.05, str(v), ha="center", va="bottom")
    fig.tight_layout()
    fig.savefig(out, dpi=300)
    plt.close(fig)
    return out


def stream_signature_probe(max_rows: int | None = None) -> pd.DataFrame:
    signature_path = PROJECT_ROOT / "computational_runs" / "result_provenance_2026-05-16" / "result_signature_summary.csv"
    if not signature_path.exists():
        return pd.DataFrame()
    rows = []
    with signature_path.open("r", encoding="utf-8-sig", errors="replace", newline="") as fh:
        reader = csv.DictReader(fh)
        for i, row in enumerate(reader):
            if max_rows is not None and i >= max_rows:
                break
            sig = row.get("parameter_signature", "")
            sample = row.get("sample_result_relative_path", "")
            if not sig:
                continue
            lower = (sig + " " + sample).lower()
            if ("cf1f=0" in lower and "cf2f=1" in lower and "cf3f=0" in lower) or "cf2f1.0" in lower:
                if any(token in lower for token in ["elasticdata", "plastic", "result", "nodedeform"]):
                    family = "separated_or_sepH" if any(t in sample for t in ["分离", "sep", "Abaqus_dataTEsepH"]) else "overall_or_unknown"
                    rows.append(
                        {
                            "matched_family_bucket": row.get("matched_family_bucket", ""),
                            "result_role": row.get("result_role", ""),
                            "matched_fieldset": row.get("matched_fieldset", ""),
                            "parameter_signature": sig[:800],
                            "result_file_count": row.get("result_file_count", ""),
                            "matched_case_count": row.get("matched_case_count", ""),
                            "sample_result_relative_path": sample,
                            "family_guess": family,
                        }
                    )
    return pd.DataFrame(rows)


def write_markdown_table(path: Path, df: pd.DataFrame, title: str, note: str = "") -> None:
    lines = [f"# {title}", ""]
    if note:
        lines.extend([note, ""])
    if df.empty:
        lines.append("No rows were found.")
    else:
        lines.append(df.to_markdown(index=False))
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def raw_solver_report(inventory: pd.DataFrame, audit: pd.DataFrame) -> Path:
    path = DOC_DIR / "P1_RAW_SOLVER_EVIDENCE_AUDIT.md"
    job2 = audit[audit["job"] == "Job-2"].iloc[0].to_dict()
    lines = [
        "# P1 Raw Solver Evidence Audit",
        "",
        f"Audit generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Scope",
        "",
        "This audit uses the available raw separated-family Abaqus working directory discovered outside the original submission archive. It is not treated as a complete 12-case reproduction package. Its evidentiary role is limited to traceability of available solver files, element formulation and one real nonlinear solver termination state.",
        "",
        "## File evidence",
        "",
        f"- Source directory: `{RAW_SEPH_DIR}`",
        f"- Files inventoried: {len(inventory)}",
        f"- Job-2 ODB size: {inventory.loc[inventory['file_name'].eq('Job-2.odb'), 'size_bytes'].iloc[0] if 'Job-2.odb' in set(inventory['file_name']) else NA} bytes",
        f"- Job-2 INP SHA-256(first 1 MB): `{inventory.loc[inventory['file_name'].eq('Job-2.inp'), 'sha256_first_1mb'].iloc[0] if 'Job-2.inp' in set(inventory['file_name']) else NA}`",
        "",
        "## Job-2 parsed findings",
        "",
        "| Item | Finding |",
        "| --- | --- |",
        f"| Element type from INP | {job2.get('element_types', NA)} |",
        f"| Riks step from INP | {job2.get('has_static_riks', NA)} |",
        f"| Contact/friction from INP | contact={job2.get('has_contact', NA)}, friction={job2.get('has_friction', NA)} |",
        f"| Plastic material block from INP | {job2.get('has_plastic_material', NA)} |",
        f"| STA completion flag | {job2.get('sta_completed_successfully', NA)} |",
        f"| STA final row | `{job2.get('sta_final_line', '')}` |",
        f"| MSG termination due to previous errors | {job2.get('analysis_terminated_due_to_errors', NA)} |",
        f"| Minimum time increment error | {job2.get('minimum_time_increment_error', NA)} |",
        f"| Total increments / cutbacks / iterations | {job2.get('total_increments', NA)} / {job2.get('cutbacks', NA)} / {job2.get('iterations', NA)} |",
        f"| Negative-eigenvalue warnings / error messages | {job2.get('negative_eigenvalue_warnings', NA)} / {job2.get('error_messages', NA)} |",
        f"| Wallclock time | {job2.get('wallclock_seconds', NA)} s |",
        f"| Result2 max RF2 from report | {job2.get('max_rf2', NA)} at X={job2.get('max_rf2_at_x', NA)} |",
        "",
        "## Manuscript implication",
        "",
        "The raw files newly verify an available S4R shell-element/Riks/contact workflow for a separated-family Abaqus run, and they provide an auditable solver-status example. The Job-2 evidence does not support re-labelling the curve-derived RF peak as an ultimate load. The parsed `.sta/.msg` files show that this run was written to the last converged increment after nonlinear convergence failure. The manuscript should therefore keep the nonlinear metric as a curve-derived reaction index and use this audit to strengthen reproducibility and boundary disclosure, not to claim capacity validation.",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def softwarex_report(df: pd.DataFrame, plot_path: Path) -> Path:
    path = DOC_DIR / "SOFTWAREX_SHEAR_BENCHMARK_AUDIT.md"
    if df.empty:
        summary = "No SoftwareX work-directory solver files were found."
    else:
        n = len(df)
        completed = int(df["riks_sta_completed_successfully"].fillna(False).sum())
        element_types = "; ".join(sorted(set(x for x in df["riks_element_types"].dropna().astype(str) if x and x != NA)))
        summary = (
            f"The audit found {n} Riks solver job records in the SoftwareX shear benchmark work directory; "
            f"{completed}/{n} had `.sta` files marked as successfully completed. "
            f"Parsed Riks INP element types: {element_types or NA}."
        )
    lines = [
        "# SoftwareX Shear Benchmark Solver Audit",
        "",
        f"Audit generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Scope",
        "",
        "This audit checks the local SoftwareX shear-code work directory as an external solver-file benchmark. It is not used as direct validation of the present P1 paired archive, because geometry, parameter design and study purpose differ. Its role is to show that the broader project ecosystem contains solver-complete Abaqus shear workflows and to define a reproducibility target for the present P1 archive.",
        "",
        "## Summary",
        "",
        summary,
        "",
        f"Status figure: `{plot_path}`",
        "",
        "## Manuscript implication",
        "",
        "The benchmark audit should be described only as a workflow contrast and reproducibility reference. It can strengthen the argument that future P1 submission should include per-case `.inp/.sta/.msg/.odb` bundles or public derived solver logs. It should not be described as independent validation of the P1 sensitivity ratios.",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def softwarex_validation_report(df: pd.DataFrame) -> Path:
    path = DOC_DIR / "SOFTWAREX_SCANDella_VALIDATION_AUDIT.md"
    lines = [
        "# SoftwareX Scandella Validation Audit",
        "",
        f"Audit generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Scope",
        "",
        "This audit parses the local SoftwareX validation summary for Scandella et al. steel-plate-girder shear specimens G1-G4. These results are evidence that a related local Abaqus shear workflow has been checked against experiments. They are not direct validation of the P1 paired shell-representation sensitivity ratios.",
        "",
    ]
    if df.empty:
        lines.append("No validation rows were parsed.")
    else:
        mean_abs = df["abs_error_percent"].mean()
        max_abs = df["abs_error_percent"].max()
        lines.extend(
            [
                f"Parsed specimens: {len(df)}.",
                f"Mean absolute percentage error: {mean_abs:.2f}%.",
                f"Maximum absolute percentage error: {max_abs:.2f}%.",
                "",
                df[["specimen", "fe_predicted_shear_kN", "experimental_shear_kN", "signed_error_percent", "reported_lpf"]].to_markdown(index=False),
                "",
                "Manuscript implication: these values may be used to justify the SoftwareX workflow as an external benchmark contrast, but the manuscript must state that the benchmark geometry/workflow differs from the P1 paired archive and is not merged into the P1 statistical evidence.",
            ]
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def deconfounded_probe_report(df: pd.DataFrame) -> Path:
    path = DOC_DIR / "DECONFOUNDED_MATCH_PROBE.md"
    lines = [
        "# Deconfounded Match Probe",
        "",
        f"Audit generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Scope",
        "",
        "This is a conservative stream probe over the large result-provenance signature table. It searches for result signatures related to the cf1f=0/cf2f=1/cf3f=0 in-plane shear baseline. The probe is not treated as a matched-interface rerun and does not create new capacity evidence.",
        "",
    ]
    if df.empty:
        lines.extend(
            [
                "No matching signature rows were found by the conservative token search.",
                "",
                "Manuscript consequence: do not claim deconfounded interface evidence. Keep matched-interface reruns as required future work.",
            ]
        )
    else:
        counts = df.groupby(["family_guess", "result_role"], dropna=False).size().reset_index(name="signature_rows")
        lines.extend(
            [
                f"Rows found: {len(df)}",
                "",
                "Grouped counts:",
                "",
                counts.to_markdown(index=False),
                "",
                "Manuscript consequence: candidate result signatures exist, but this probe is insufficient to establish a deconfounded same-interface comparison. The safe use is to define a targeted rerun/indexing plan, not to add a new result claim.",
            ]
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def copy_final_figures(plot_path: Path) -> None:
    final_dir = SUP_ROOT / "final_figures"
    final_dir.mkdir(parents=True, exist_ok=True)
    if plot_path.exists():
        shutil.copy2(plot_path, final_dir / plot_path.name)


def build_traceability(raw_report: Path, softwarex_report_path: Path, probe_report: Path) -> Path:
    trace = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "new_evidence_items": [
            {
                "id": "E-SOLVER-RAW-SEPH-JOB2",
                "description": "Available separated-family raw Abaqus Job-2 INP/STA/MSG/ODB files; verifies S4R/Riks/contact setup and one nonlinear termination state.",
                "source_directory": str(RAW_SEPH_DIR),
                "derived_outputs": [
                    str(OUT_DIR / "p1_raw_solver_file_inventory.csv"),
                    str(OUT_DIR / "p1_raw_job_solver_audit.csv"),
                    str(raw_report),
                ],
                "allowed_claim": "Element/method traceability improved for available raw separated-family solver files.",
                "forbidden_claim": "This does not validate all P1 cases and does not establish ultimate capacity.",
            },
            {
                "id": "E-SOLVER-SOFTWAREX-SHEAR",
                "description": "Local SoftwareX shear workflow solver-file audit used as reproducibility benchmark contrast.",
                "source_directory": str(SOFTWAREX_WORK),
                "derived_outputs": [
                    str(OUT_DIR / "softwarex_shear_solver_benchmark_audit.csv"),
                    str(softwarex_report_path),
                ],
                "allowed_claim": "A solver-complete reference workflow exists in the broader project ecosystem.",
                "forbidden_claim": "This is not direct external validation of the P1 sensitivity ratios.",
            },
            {
                "id": "E-SOFTWAREX-SCANDELLA-VALIDATION",
                "description": "Local SoftwareX validation summary for Scandella et al. G1-G4 shear specimens.",
                "derived_outputs": [
                    str(OUT_DIR / "softwarex_scandella_validation_summary.csv"),
                    str(DOC_DIR / "SOFTWAREX_SCANDella_VALIDATION_AUDIT.md"),
                ],
                "allowed_claim": "A related local shear workflow has reported experiment-comparison errors of 0.6-7.1% across four Scandella specimens.",
                "forbidden_claim": "These results do not directly validate the P1 shell-representation sensitivity ratios.",
            },
            {
                "id": "E-DECONFOUNDED-PROBE",
                "description": "Conservative result-provenance token probe for candidate in-plane shear signatures.",
                "derived_outputs": [
                    str(OUT_DIR / "deconfounded_match_probe_rows.csv"),
                    str(probe_report),
                ],
                "allowed_claim": "Candidate indexing/provenance paths can guide future matched reruns.",
                "forbidden_claim": "No new deconfounded matched-interface result is claimed.",
            },
        ],
    }
    out = DOC_DIR / "TRACEABILITY_V8_SUPPLEMENT.json"
    out.write_text(json.dumps(trace, indent=2, ensure_ascii=False), encoding="utf-8")
    return out


def revise_v8(raw_audit: pd.DataFrame, softwarex_df: pd.DataFrame) -> Path:
    original = read_text(V7_MANUSCRIPT)
    if not original:
        raise FileNotFoundError(V7_MANUSCRIPT)

    job2 = raw_audit[raw_audit["job"] == "Job-2"].iloc[0].to_dict()
    sw_n = len(softwarex_df)
    sw_completed = int(softwarex_df["riks_sta_completed_successfully"].fillna(False).sum()) if sw_n else 0
    validation_path = OUT_DIR / "softwarex_scandella_validation_summary.csv"
    validation_df = pd.read_csv(validation_path) if validation_path.exists() else pd.DataFrame()
    if validation_df.empty:
        validation_sentence = "A local experimental validation summary was not parsed for this v8 build."
    else:
        validation_sentence = (
            f"The same SoftwareX validation records report {len(validation_df)} Scandella et al. shear specimens with "
            f"absolute percentage errors ranging from {validation_df['abs_error_percent'].min():.1f}% to "
            f"{validation_df['abs_error_percent'].max():.1f}% and a mean absolute error of "
            f"{validation_df['abs_error_percent'].mean():.2f}%."
        )

    odb_size = job2.get("odb_size_bytes", NA)
    try:
        odb_size_text = f"{int(odb_size):,}"
    except (TypeError, ValueError):
        odb_size_text = str(odb_size)

    insert_section = f"""
### 2.5 Raw solver-output audit and benchmark contrast

After the v7 audit, an additional separated-family Abaqus working directory was located outside the original submission archive. The available raw files include `Job-2.inp`, `Job-2.sta`, `Job-2.msg` and a `Job-2.odb` file of {odb_size_text} bytes. The parsed input file verifies S4R shell elements, shell sections, a geometrically nonlinear Static Riks step, boundary conditions, contact and friction definitions, and a plastic material block. This newly supports element- and solver-setting traceability for the available separated-family raw run.

The same audit also sharpens the interpretation boundary. The parsed `Job-2.sta` file is not marked as successfully completed; the final status line reports that the analysis was not completed. The parsed message file reports termination due to previous errors, a minimum-time-increment error, {job2.get('total_increments', 'not_available')} increments, {job2.get('cutbacks', 'not_available')} cutbacks, {job2.get('iterations', 'not_available')} iterations and {job2.get('negative_eigenvalue_warnings', 'not_available')} negative-eigenvalue warnings. Therefore, this raw solver evidence strengthens traceability but does not permit the nonlinear RF peak to be interpreted as an ultimate load or design resistance. It supports the present choice to report peak_abs_rf2 only as a curve-derived reaction index.

For external workflow contrast, the local SoftwareX shear-code work directory was also audited. The audit found {sw_n} Riks shear benchmark job records, of which {sw_completed} had `.sta` files marked as successfully completed. {validation_sentence} These files are not used as direct validation of the present paired sensitivity ratios because they belong to a different published software-validation workflow. They are used only to define the reproducibility standard that the present P1 archive should meet in a stronger resubmission: per-case input files, solver-status files, output databases or derived solver logs, and a short runbook linking each manuscript number to its source file.
"""

    if "### 2.5 Raw solver-output audit and benchmark contrast" not in original:
        original = original.replace("## 3. Results", insert_section.strip() + "\n\n## 3. Results")

    # Update the method-facts paragraph now that an explicit raw INP was found.
    original = original.replace(
        "The archived generator files available for this audit do not expose a verified explicit element-type assignment, and the submission archive does not contain the Job-2 ODB files needed to reconstruct mesh-level convergence or failure-mode histories. These missing items are treated as publication boundaries rather than inferred retrospectively.",
        "The original submission archive did not contain the Job-2 ODB/log files needed to reconstruct mesh-level convergence or failure-mode histories. A subsequently located separated-family raw working directory verifies an S4R shell-element Riks/contact setup for an available raw run, but it is not a complete matched-case solver package for the paired study. The remaining missing per-case ODB/log coverage is treated as a publication boundary rather than inferred retrospectively.",
    )

    # Strengthen data availability.
    original = original.replace(
        "The original submission archive does not contain the corresponding Job-2 solver logs or ODB files.",
        "The original submission archive does not contain the corresponding per-case Job-2 solver logs or ODB files. A separately located separated-family working directory contains raw Job-2 solver files and is audited as traceability evidence, but it is not a full matched archive for all paired cases.",
    )

    # Adjust final probability language if present or append an editorial note.
    editorial_note = f"""

## Editorial evidence note for v8

The v8 supplement improves auditability by adding raw Abaqus solver-file evidence and a solver-complete SoftwareX benchmark contrast. It does not convert the study into an experimental validation paper or a design-capacity paper. A realistic submission strategy remains to present the manuscript as a bounded FE modelling-sensitivity/technical note. The most acceptance-limiting missing evidence is still a matched per-case solver package or rerun matrix that independently varies shell representation and interface assumptions.
"""
    if "## Editorial evidence note for v8" not in original:
        original = original.rstrip() + "\n" + editorial_note

    ai_declaration = """## Declaration of generative AI and AI-assisted technologies in the writing process

During the preparation of this manuscript draft, OpenAI ChatGPT/Codex was used to support manuscript organization, consistency checking, reproducibility-audit scripting and language editing. The scientific interpretation, file verification, numerical reporting and final submission responsibility remain with the human authors.
"""
    if "## Declaration of generative AI and AI-assisted technologies in the writing process" not in original:
        original = original.replace("## References", ai_declaration.strip() + "\n\n## References")
    if "Scandella C, Neuenschwander M, Mosalam KM, Knobloch M, Fontana M." not in original:
        original = original.replace(
            "[5] Daley AJ, Davis DB, White DW. Shear strength of unstiffened steel plate girders. In: Structural Stability Research Council Annual Stability Conference 2016, SSRC 2016. 2016. pp. 132-147.",
            "[5] Daley AJ, Davis DB, White DW. Shear strength of unstiffened steel plate girders. In: Structural Stability Research Council Annual Stability Conference 2016, SSRC 2016. 2016. pp. 132-147.\n\n[6] Scandella C, Neuenschwander M, Mosalam KM, Knobloch M, Fontana M. Structural behavior of steel-plate girders in shear: Experimental study and review of current design principles. Journal of Structural Engineering. 2020;146(11):04020243. doi:10.1061/(ASCE)ST.1943-541X.0002804.",
        )

    out = MANUSCRIPT_DIR / "manuscript_v8_high_probability_supplemented.md"
    out.write_text(original, encoding="utf-8")
    return out


def add_paragraph_with_style(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)
    return p


def markdown_to_docx(md_path: Path, docx_path: Path) -> Path:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for level in range(1, 4):
        style = styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(16 - level)
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    for raw in read_text(md_path).splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.add_run(line[2:].strip())
        elif line.startswith("|"):
            add_paragraph_with_style(doc, line)
        else:
            add_paragraph_with_style(doc, line)
    doc.save(docx_path)
    return docx_path


def write_highlights() -> Path:
    highlights = [
        "Quantifies first-eigenvalue sensitivity in paired H-section shell models.",
        "Bounds nonlinear RF-U response as curve-derived, not design resistance.",
        "Verifies S4R/Riks/contact settings from newly audited Abaqus files.",
        "Maps missing solver evidence required before capacity claims are defensible.",
    ]
    out = PACKAGE_DIR / "Highlights_v8.txt"
    out.write_text("\n".join(highlights) + "\n", encoding="utf-8")
    return out


def write_cover_letter() -> Path:
    out = PACKAGE_DIR / "CoverLetter_JCSR_v8.docx"
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    paragraphs = [
        "Subject: Submission of a manuscript to Journal of Constructional Steel Research",
        "Dear Editor, we submit the manuscript entitled \"Representation sensitivity in shell finite-element modelling of H-section steel members under in-plane shear\" for consideration as a bounded finite-element modelling-sensitivity contribution.",
        "The manuscript reports an audited paired Abaqus archive comparing overall/equivalent and separated H-section shell model families under an in-plane shear baseline. The strongest supported result is elastic: 11 paired cases with positive Mode 1 eigenvalues give a separated-to-overall mean ratio of 0.961 with a 95% bootstrap confidence interval of 0.944-0.978. The nonlinear RF-U evidence is deliberately reported as a curve-derived reaction index, not as ultimate shear resistance.",
        "We believe the manuscript fits the journal because it addresses a practical modelling issue in structural steel finite-element analysis: how shell representation, contact/interface assumptions and post-processing can alter response metrics and lead to over-interpretation if solver evidence is incomplete.",
        "The work is original, is not under consideration elsewhere, and all authors must confirm authorship, competing-interest, funding, data-availability and AI-assistance declarations before submission. Suggested reviewers may be entered in the editorial system by the authors; no reviewer names are invented in this draft package.",
    ]
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(8)
    doc.save(out)
    return out


def write_submission_readme(package_paths: dict[str, Path]) -> Path:
    out = PACKAGE_DIR / "00-README-presubmission-v8.md"
    lines = [
        "# JCSR v8 投稿前必读",
        "",
        "## 目标期刊与定位",
        "",
        "- 目标期刊：Journal of Constructional Steel Research (Elsevier).",
        "- 推荐投稿定位：bounded FE modelling-sensitivity / technical note style, not design-capacity paper.",
        "- 2026-05-17 核验：JCSR Guide for Authors 要求数据可用性声明；参考文献正文用方括号数字引用；Elsevier 要求使用生成式 AI 时在参考文献前放声明。",
        "",
        "## 相比 v7 的实质增量",
        "",
        "- 新增可审计原始 Abaqus separated-family `Job-2.inp/.sta/.msg/.odb` 文件审计。",
        "- 从 raw INP 验证 `S4R`、Static Riks、contact/friction 和 plastic material block。",
        "- 从 raw STA/MSG 解析出 Job-2 未成功完成、108 increments、20 cutbacks、447 iterations、2 error messages。",
        "- 新增 SoftwareX shear workflow solver-status benchmark contrast：4/4 Riks records completed。",
        "- 新增 v8 正文 2.5 节，将 solver traceability 与 capacity validation 严格分开。",
        "",
        "## 作者必须人工确认的 5 项",
        "",
        "1. 作者列表、通讯作者地址、电话、邮箱和 ORCID。",
        "2. Funding statement；如果无经费，确认是否使用 Elsevier 推荐的 no-funding statement。",
        "3. Competing interests declaration。",
        "4. 原始数据/派生表是否可公开；如不能公开，需要说明不可公开原因。",
        "5. AI declaration 的具体工具名称、用途和作者审阅责任表述。",
        "",
        "## EM 系统建议上传顺序",
        "",
        "1. Manuscript DOCX。",
        "2. Highlights TXT。",
        "3. Cover letter DOCX。",
        "4. Supplementary audit tables and reports, if allowed。",
        "5. Figures/tables and derived CSV evidence package。",
        "",
        "## 本包文件",
        "",
    ]
    for label, path in package_paths.items():
        lines.append(f"- {label}: `{path}`")
    lines.extend(
        [
            "",
            "## 仍然不应主张",
            "",
            "- 不应写 ultimate load / collapse load / design resistance。",
            "- 不应写 separated shell model is generally more accurate/safer。",
            "- 不应写 SoftwareX benchmark directly validates P1 ratios。",
            "- 不应写 deconfounded interface effect has been proven。",
        ]
    )
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def prohibited_scan(md_path: Path) -> Path:
    terms = [
        "fundamental shift",
        "first-of-its-kind",
        "ground-breaking",
        "direct SEM",
        "direct XRD",
        "complete framework",
        "comprehensive framework",
        "wasted",
        "inefficient",
        "significant improvement",
        "significantly improves",
        "proved",
        "proves",
    ]
    text = read_text(md_path)
    rows = []
    for term in terms:
        pattern = r"\b" + re.escape(term) + r"\b"
        for m in re.finditer(pattern, text, flags=re.IGNORECASE):
            start = max(0, m.start() - 80)
            end = min(len(text), m.end() + 80)
            rows.append({"term": term, "context": text[start:end].replace("\n", " ")})
    out_csv = PACKAGE_DIR / "PROHIBITED_SCAN_v8.csv"
    pd.DataFrame(rows, columns=["term", "context"]).to_csv(out_csv, index=False, encoding="utf-8-sig")
    out_md = PACKAGE_DIR / "PROHIBITED_SCAN_v8.md"
    if rows:
        body = pd.DataFrame(rows).to_markdown(index=False)
    else:
        body = "No prohibited phrases from the configured scan list were found."
    out_md.write_text("# Prohibited Phrase Scan v8\n\n" + body + "\n", encoding="utf-8")
    return out_md


def jcsr_instruction_check() -> Path:
    out = DOC_DIR / "JCSR_AUTHOR_INSTRUCTIONS_CHECK_2026-05-17.md"
    lines = [
        "# JCSR Author Instructions Check",
        "",
        "Checked date: 2026-05-17.",
        "",
        "Official sources consulted:",
        "",
        "- Journal of Constructional Steel Research Guide for Authors, ScienceDirect/Elsevier: https://www.sciencedirect.com/journal/journal-of-constructional-steel-research/publish/guide-for-authors",
        "- Elsevier policy on AI and AI-assisted technologies in writing: https://www.elsevier.com/en-gb/about/policies-and-standards/the-use-of-generative-ai-and-ai-assisted-technologies-in-writing-for-elsevier",
        "",
        "Submission-relevant points recorded for this package:",
        "",
        "- JCSR considers high-standard original papers on theoretical and experimental research in steel and metal structures.",
        "- JCSR uses single-anonymized peer review and initial editorial suitability screening.",
        "- JCSR requires a data availability statement at submission.",
        "- JCSR reference style uses numbered references in square brackets in order of appearance.",
        "- Elsevier requires disclosure when generative AI tools were used in manuscript preparation; the statement should appear before the references.",
        "- Elsevier does not permit generative AI or AI-assisted tools to create or alter submitted scientific images, except where such use is part of the reproducible research method.",
    ]
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def build_submission_package(v8_md: Path) -> dict[str, Path]:
    manuscript_docx = markdown_to_docx(v8_md, PACKAGE_DIR / "Manuscript_JCSR_v8_high_probability_supplemented.docx")
    highlights = write_highlights()
    cover = write_cover_letter()
    scan = prohibited_scan(v8_md)
    instructions = jcsr_instruction_check()
    package_paths = {
        "manuscript": manuscript_docx,
        "highlights": highlights,
        "cover letter": cover,
        "prohibited scan": scan,
        "JCSR instruction check": instructions,
    }
    readme = write_submission_readme(package_paths)
    package_paths["submission README"] = readme
    return package_paths


def gate_report(raw_audit: pd.DataFrame, softwarex_df: pd.DataFrame, outputs: dict[str, Path]) -> Path:
    job2 = raw_audit[raw_audit["job"] == "Job-2"].iloc[0]
    sw_n = len(softwarex_df)
    sw_completed = int(softwarex_df["riks_sta_completed_successfully"].fillna(False).sum()) if sw_n else 0
    validation_path = OUT_DIR / "softwarex_scandella_validation_summary.csv"
    validation_df = pd.read_csv(validation_path) if validation_path.exists() else pd.DataFrame()
    validation_line = ""
    if not validation_df.empty:
        validation_line = f"- Parsed SoftwareX/Scandella validation summary: n={len(validation_df)}, mean absolute error={validation_df['abs_error_percent'].mean():.2f}%, range={validation_df['abs_error_percent'].min():.1f}-{validation_df['abs_error_percent'].max():.1f}%."
    lines = [
        "# V8 High-Probability Supplement Gate Report",
        "",
        "GATE = EVIDENCE IMPROVED, CAPACITY CLAIM STILL BLOCKED",
        "",
        "## What was materially added",
        "",
        "- Parsed an available raw separated-family Abaqus working directory containing `Job-2.inp/.sta/.msg/.odb`.",
        f"- Verified from raw INP: element type `{job2.get('element_types', NA)}`, Static Riks={job2.get('has_static_riks', NA)}, contact={job2.get('has_contact', NA)}, friction={job2.get('has_friction', NA)}.",
        f"- Parsed raw solver status: completed={job2.get('sta_completed_successfully', NA)}, increments={job2.get('total_increments', NA)}, cutbacks={job2.get('cutbacks', NA)}, iterations={job2.get('iterations', NA)}, error messages={job2.get('error_messages', NA)}.",
        f"- Audited SoftwareX shear benchmark solver files: {sw_completed}/{sw_n} Riks records marked completed.",
        validation_line,
        "- Added v8 manuscript section separating solver traceability evidence from validation/capacity evidence.",
        "",
        "## What remains blocked",
        "",
        "- The available raw separated-family directory is not a full 12-case matched solver archive.",
        "- The parsed Job-2 status is not successfully completed; it reinforces last-converged/curve-index wording.",
        "- No matched-interface reruns were completed, so representation and interface/contact effects remain coupled.",
        "- SoftwareX benchmark files are a reproducibility contrast, not direct validation of the P1 ratios.",
        "",
        "## Probability update",
        "",
        "- JCSR Original Article: approximately 25-35% after v8, mainly because auditability and element-type traceability are stronger.",
        "- JCSR short technical note / modelling note: approximately 35-45% if the manuscript is submitted with the bounded title, v8 solver audit and all derived tables.",
        "- A genuinely high-probability JCSR package still requires matched per-case solver logs/ODBs or reruns and preferably a deconfounded interface study.",
        "",
        "## Output files",
        "",
    ]
    for label, path in outputs.items():
        lines.append(f"- {label}: `{path}`")
    out = DOC_DIR / "V8_HIGH_PROBABILITY_SUPPLEMENT_GATE_REPORT.md"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def main() -> None:
    ensure_dirs()
    inventory, raw_audit = raw_solver_audit()
    inventory.to_csv(OUT_DIR / "p1_raw_solver_file_inventory.csv", index=False, encoding="utf-8-sig")
    raw_audit.to_csv(OUT_DIR / "p1_raw_job_solver_audit.csv", index=False, encoding="utf-8-sig")

    softwarex_df = softwarex_solver_audit()
    softwarex_df.to_csv(OUT_DIR / "softwarex_shear_solver_benchmark_audit.csv", index=False, encoding="utf-8-sig")
    validation_df = softwarex_validation_audit()
    validation_df.to_csv(OUT_DIR / "softwarex_scandella_validation_summary.csv", index=False, encoding="utf-8-sig")
    plot_path = make_softwarex_plot(softwarex_df)
    copy_final_figures(plot_path)

    probe_df = stream_signature_probe()
    probe_df.to_csv(OUT_DIR / "deconfounded_match_probe_rows.csv", index=False, encoding="utf-8-sig")

    raw_report = raw_solver_report(inventory, raw_audit)
    softwarex_report_path = softwarex_report(softwarex_df, plot_path)
    validation_report_path = softwarex_validation_report(validation_df)
    probe_report = deconfounded_probe_report(probe_df)
    trace_path = build_traceability(raw_report, softwarex_report_path, probe_report)
    v8_path = revise_v8(raw_audit, softwarex_df)
    package_paths = build_submission_package(v8_path)

    outputs = {
        "raw solver inventory": OUT_DIR / "p1_raw_solver_file_inventory.csv",
        "raw solver audit": OUT_DIR / "p1_raw_job_solver_audit.csv",
        "SoftwareX benchmark audit": OUT_DIR / "softwarex_shear_solver_benchmark_audit.csv",
        "SoftwareX Scandella validation summary": OUT_DIR / "softwarex_scandella_validation_summary.csv",
        "deconfounded probe rows": OUT_DIR / "deconfounded_match_probe_rows.csv",
        "raw solver report": raw_report,
        "SoftwareX report": softwarex_report_path,
        "SoftwareX Scandella validation report": validation_report_path,
        "deconfounded probe report": probe_report,
        "traceability": trace_path,
        "v8 manuscript": v8_path,
        "status figure": plot_path,
        **{f"package {k}": v for k, v in package_paths.items()},
    }
    gate = gate_report(raw_audit, softwarex_df, outputs)
    outputs["gate report"] = gate

    print("V8 supplement build completed.")
    for label, path in outputs.items():
        print(f"{label}: {path}")


if __name__ == "__main__":
    main()
